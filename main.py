import docx

document = docx.Document("PDC.docx")
course_title = ""
description = ""
course_goal = ""
impact = ""
staffresource = ""
isDesc = False
descCounter = 0
isCourseGoal = False
courseGoalCounter = 0
isImpact = False
impactCounter = 0
isStaffResource = False
staffResourceCounter = 0
isReliance = False
relianceCounter = 0
reliance = ""
isImplication = 0
implicationCounter = 0
implication = ""
isAnticipatedResources = False
anticipatedResourcesCounter = 0
anticipated_resources = ""
isPlannedRelocation = False
plannedRelocationCounter = 0
planned_resource = ""
pre_requisites = ""
co_requisites = ""
anti_requisites = ""
cross_list = ""
required_course = ""
replacing_old_course = ""
#

title_of_program = ""
school = ""
faculty = ""
proposed_changes = ""
credit_weight = 0
total_contact_hours = 0
in_class_hrs = 0
e_learning_hrs = 0
distance_hrs = 0
other_hrs = 0
lecture_hrs = 0
lab_hrs = 0
online_hrs = 0
coop_hrs = 0
willGetCredit = False
course_learning_outcomes = []
demand_for_course = []
student_workload = {}
additional_resource = {}
additional_institute_resource = {}
form_history = []
# looping to paragraph
for paragaph in document.paragraphs:
    text = paragaph.text
    #description
    if isDesc and descCounter == 0:
        descCounter += 1
    if isDesc and descCounter == 1:
        isDesc = False
        description = text
    # for course goal
    if isCourseGoal and courseGoalCounter == 0:
        courseGoalCounter += 1
    if isCourseGoal and courseGoalCounter == 1:
        isCourseGoal = False
        course_goal = text
    # for impact counter
    if isImpact and impactCounter == 0:
        impactCounter += 1
    if isImpact and impactCounter == 1:
        isImpact = False
        impact = text
    # for staff resource
    if isStaffResource and staffResourceCounter == 0:
        staffResourceCounter += 1
    if isStaffResource and staffResourceCounter == 1:
        isStaffResource = False
        staffresource = text
    # reliance
    if isReliance and relianceCounter == 0:
        relianceCounter += 1
    if isReliance and relianceCounter == 1:
        isReliance = False
        reliance = text
    # implication
    if isImplication and implicationCounter == 0:
        implicationCounter += 1
    if isImplication and implicationCounter == 1:
        isImplication = False
        implication = text
    # anticipated resources
    if isAnticipatedResources and anticipatedResourcesCounter == 0:
        anticipatedResourcesCounter += 1
    if isAnticipatedResources and anticipatedResourcesCounter == 1:
        isAnticipatedResources = False
        anticipated_resources = text
    #planned reloaction
    if isPlannedRelocation and plannedRelocationCounter == 0:
        plannedRelocationCounter += 1
    if isPlannedRelocation and plannedRelocationCounter == 1:
        isPlannedRelocation = False
        planned_resource = text
    # checking for the data
    if "Course # and Title:" in text:
        # splitting the string after : to get the course number and title
        split_string = text.split(":")
        course_title = split_string[len(split_string) - 1]
        course_title = course_title.strip()
    elif "A.1 Calendar Description" in text:
        # using a flag to know that the line after this line will be the calendar description
        isDesc = True
    elif "B.1	Course Goal(s)" in text:
        isCourseGoal = True
    elif "B.4.1  Impact of New Course on Enrolment in Existing Courses" in text:
        isImpact = True
    elif "C.1	Available Faculty Expertise and Staff Resources (QAF sections 2.1.7, 2.1.8, 2.1.9 and 2.1.10)" in text:
        isStaffResource = True
    elif "C.1.1 Extent of Reliance on Adjunct, Limited-term, and Sessional Faculty in Delivering the Revised Program" in text:
        isReliance = True
    elif "C.2	Resource Implications for Other Campus Units (Ministry sections 3 and 4)" in text:
        isImplication = True
    elif "C.3	Anticipated New Resources (QAF sections 2.1.7, 2.1.8 and 2.1.9; Ministry section 4)" in text:
        isAnticipatedResources = True
    elif "C.4	Planned Reallocation of Resources and Cost-Savings (QAF section 2.1.7 and 2.1.9; Ministry section 4)" in text:
        isPlannedRelocation = True

# variables
print(course_title)
print(description)
print(course_goal)
print(impact)
print(staffresource)
print(reliance)
print(implication)
print(anticipated_resources)
print(planned_resource)

# fetching all tables
cnt = 0
for table in document.tables:
    if cnt == 0:
        # the title and department table
        title_of_program = table.cell(0, 1).text
        school = table.cell(1, 1).text
        faculty = table.cell(2, 1).text
    elif cnt == 1:
        # the proposed changes table
        proposed_changes = table.cell(0, 1).text
    elif cnt == 5:
        # credit table
        credit_weight = table.cell(2, 0).text
        total_contact_hours = table.cell(2, 1).text
        in_class_hrs = table.cell(2, 2).text
        e_learning_hrs = table.cell(2, 3).text
        distance_hrs = table.cell(2, 4).text
        other_hrs = table.cell(2, 5).text
        lecture_hrs = table.cell(2, 6).text
        lab_hrs = table.cell(2, 7).text
        online_hrs = table.cell(2, 8).text
        coop_hrs = table.cell(2, 9).text
    elif cnt == 6:
        # prequisities table
        pre_requisites = table.cell(1, 0).text
        co_requisites = table.cell(1, 1).text
        anti_requisites = table.cell(1, 2).text
        cross_list = table.cell(1, 3).text
        required_course = table.cell(1, 4).text
        replacing_old_course = table.cell(1, 5).text
    elif cnt == 8:
        #	Will students be able to obtain credit for the new course and the course(s) that it is replacing?
        temp = table.cell(0, 1).text
        if temp.lower() == 'yes' or temp.lower() == 'true':
            willGetCredit = True
        elif temp.lower() == 'no' or temp.lower() == 'false':
            willGetCredit = False
    elif cnt == 12:
        # Course Learning Outcomes
        for i in range(0, 9):
            temp_outcome = table.cell(i + 2, 0).text
            # splitting the text from .  since we dont want the leading alphabets from the form
            splitted_text = temp_outcome.split(".")
            course_learning_outcomes.insert(i, splitted_text[1])
    elif cnt == 14:
        # demand for course
        for i in range(0, 5):
            demand_for_course.insert(i, table.cell(1, i + 1).text)
    elif cnt == 17:
        # student workload
        student_workload["lectures"] = table.cell(1, 0).text
        student_workload["tutorials"] = table.cell(2, 0).text
        student_workload["labs"] = table.cell(3, 0).text
        student_workload["practical_exp"] = table.cell(4, 0).text
        student_workload["independent_study"] = table.cell(5, 0).text
        student_workload["reading"] = table.cell(6, 0).text
        student_workload["work"] = table.cell(7, 0).text
        student_workload["group"] = table.cell(8, 0).text
        student_workload["tests"] = table.cell(9, 0).text
        student_workload["other"] = table.cell(10, 0).text
        student_workload["other_reason"] = table.cell(10, 2).text
        student_workload["comparison"] = table.cell(11, 3).text
    elif cnt == 24:
        # additional resources required
        additional_resource["faculty"] = table.cell(0, 1).text
        additional_resource["staff"] = table.cell(1, 1).text
        additional_resource["gata"] = table.cell(2, 1).text
    elif cnt == 26:
        # additional institute resource
        additional_institute_resource["library"] = table.cell(0, 1).text
        additional_institute_resource["taeching"] = table.cell(1, 1).text
        additional_institute_resource["studentsupport"] = table.cell(2, 1).text
        additional_institute_resource["space"] = table.cell(3, 1).text
        additional_institute_resource["equipment"] = table.cell(4, 1).text
    elif cnt == 27:
        # history form
        count = 0
        for row in table.rows:
            current_history = {}
            if count != 0:
                current_history["dom"] = table.cell(count, 0).text
                current_history["approval_body"] = table.cell(count, 1).text
                current_history["reason_for_modification"] = table.cell(count, 2).text
                form_history.append(current_history)
            count += 1

    cnt += 1

print(title_of_program)
print(school)
print(faculty)
print(credit_weight)
print(total_contact_hours)
print(in_class_hrs)
print(e_learning_hrs)
print(distance_hrs)
print(other_hrs)
print(lecture_hrs)
print(lab_hrs)
print(online_hrs)
print(coop_hrs)
print(pre_requisites)
print(co_requisites)
print(anti_requisites)
print(cross_list)
print(required_course)
print(replacing_old_course)
print(willGetCredit)
print(course_learning_outcomes)
print(demand_for_course)
print(student_workload)
print(additional_resource)
print(additional_institute_resource)
print(form_history)
